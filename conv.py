import os
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

#前処理1

# inputフォルダからCSVファイルを読み込む
csv_filename = os.path.join('input', 'output_hazardinfo.csv')
data = pd.read_csv(csv_filename, encoding='shift-jis')


# 1行目の2列目以降を地点名として取得
point_names = data.columns[1:].tolist()

# outputフォルダにハザード情報調査告書を地点数だけ複製して出力
template_docx = os.path.join('input', 'ハザード情報調査告書フォーマット.docx')
output_folder = 'output'

if not os.path.exists(output_folder):
    os.makedirs(output_folder)

for point_name in point_names:
    output_filename = os.path.join(output_folder, f'{point_name}.docx')
    doc = Document(template_docx)
    doc.save(output_filename)
    print(f'Created {output_filename}')


#前処理2

# 入力ファイルのパス
input_file_path = 'input/output_hazardinfo.csv'

# CSVファイルを読み込む
df = pd.read_csv(input_file_path, encoding='shift-jis')

# 地点名の列を特定（2列目から右をすべて）
location_columns = df.columns[1:]

# 地点ごとにファイルを出力
for location_column in location_columns:
    location_name = location_column.strip()  # 空白を削除して地点名を取得
    output_file_name = os.path.join('input', f'{location_name}.csv')
    location_data = df[[location_column]]  # 地点のデータを抽出
    location_data.to_csv(output_file_name, index=False, encoding='shift-jis')
    print(f'{output_file_name} にデータを出力しました.')

# ファイル名のリストを作成
file_names = [f'{location_column.strip()}.csv' for location_column in location_columns]

# ファイル名に ".docs" を追加したリストを作成
file_names_with_docs = [f'{location_column.strip()}.docx' for location_column in location_columns]

# ファイル名のリストをデータフレームに追加
file_list_df = pd.DataFrame({'FileNames': file_names})
file_list_df['FileNamesWithDocs'] = file_names_with_docs

# ファイル名のリストを "拠点一覧.csv" として出力
file_list_path = os.path.join('input', '入力ファイル一覧.csv')
file_list_df.to_csv(file_list_path, index=False, encoding='shift-jis')
print(f'{file_list_path} に拠点一覧を出力しました.')


#本処理######################################################################################################################

# inputフォルダとoutputフォルダのパスを設定
input_folder = 'input'
output_folder = 'output'

# CSVファイルを読み込む
csv_file_path = os.path.join(input_folder, '入力ファイル一覧.csv')
df = pd.read_csv(csv_file_path)

# データを辞書に格納 {気象庁データ名: 対応する報告書データ名}
data_mapping = dict(zip(df.iloc[:, 0], df.iloc[:, 1]))

# すべての気象庁データ名に対して処理を実行
for meteorological_data_name, report_data_name in data_mapping.items():
    # 気象庁データのファイルパス
    meteorological_data_file = os.path.join(input_folder, f'{meteorological_data_name}')
    
    # 対応する報告書データのファイルパス
    report_data_file = os.path.join(output_folder, f'{report_data_name}')
    
    # 気象庁データを読み込む
    meteorological_df = pd.read_csv(meteorological_data_file, encoding='shift-jis')
    
    # 対応する報告書データを読み込む
    doc = Document(report_data_file)
###############################################################################################################################################
###############################################################################################################################################

#基準風速
    # 気象庁データの22行目を読み込み（基準風速）
    kijun_wind_speed = str(meteorological_df.iloc[20, 0])

    # 対応する報告書データの2つ目の表の13行目4列目のテキストを置き換える
    table = doc.tables[2]  # 2つ目の表を取得
    cell = table.cell(12, 3)  # 13行目4列目のセルを取得

    # 既存のテキストを置換
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""  # テキストをクリア

    run1 = cell.paragraphs[0].add_run(str(kijun_wind_speed+"m/sec"))
    run1.font.name = "游明朝"
    run1.font.size = Pt(10)
    run1.font.underline = False
    run1.font.bold = False  

    # 対応する報告書データの8つ目の表の1行目3列目のテキストを置き換える
    table = doc.tables[7]  # 8つ目の表を取得
    cell = table.cell(0, 1)  # 1行目2列目のセルを取得

    # 既存のテキストを置換
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""  # テキストをクリア

    run1 = cell.paragraphs[0].add_run(str(kijun_wind_speed+"m/sec"))
    run1.font.name = "游明朝"
    run1.font.size = Pt(10)
    run1.font.underline = False
    run1.font.bold = False   

    # 気象庁データの20行目を読み込み（基準風速市町村）
    kijun_place_wind_speed = str(meteorological_df.iloc[18, 0])

    # 対応する報告書データの2つ目の表の1行目3列目のテキストを置き換える
    table = doc.tables[2]  # 2つ目の表を取得
    cell = table.cell(12, 5)  # 12行目5列目のセルを取得

    # 既存のテキストを置換
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""  # テキストをクリア

    run1 = cell.paragraphs[0].add_run(str("("+kijun_place_wind_speed+")"))
    run1.font.name = "游明朝"
    run1.font.size = Pt(10)
    run1.font.underline = False
    run1.font.bold = False  

    # 対応する報告書データの8つ目の表の1行目3列目のテキストを置き換える
    table = doc.tables[7]  # 8つ目の表を取得
    cell = table.cell(0, 2)  # 1行目3列目のセルを取得

    # 既存のテキストを置換
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""  # テキストをクリア

    run1 = cell.paragraphs[0].add_run(str("("+kijun_place_wind_speed+")"))
    run1.font.name = "游明朝"
    run1.font.size = Pt(10)
    run1.font.underline = False
    run1.font.bold = False    
###############################################################################################################################################
###############################################################################################################################################

#採用データ(データ更新時に書き換え!!!!!!)

    # 対応する報告書データの8つ目の表の2行目3列目のテキストを置き換える
    table = doc.tables[8]  # 8つ目の表を取得
    cell = table.cell(2, 2)  # 2行目2列目のセルを取得

    # 既存のテキストを置換
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""  # テキストをクリア

    run1 = cell.paragraphs[0].add_run(str("1990-2021"))
    run1.font.name = "游明朝"
    run1.font.size = Pt(10)
    run1.font.underline = False
    run1.font.bold = False    

    # 対応する報告書データの8つ目の表の2行目3列目のテキストを置き換える
    table = doc.tables[8]  # 8つ目の表を取得
    cell = table.cell(3, 2)  # 3行目2列目のセルを取得

    # 既存のテキストを置換
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""  # テキストをクリア

    run1 = cell.paragraphs[0].add_run(str("1990-2021"))
    run1.font.name = "游明朝"
    run1.font.size = Pt(10)
    run1.font.underline = False
    run1.font.bold = False    

    # 対応する報告書データの9つ目の表の2行目3列目のテキストを置き換える
    table = doc.tables[9]  # 8つ目の表を取得
    cell = table.cell(2, 2)  # 3行目3列目のセルを取得

    # 既存のテキストを置換
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""  # テキストをクリア

    run1 = cell.paragraphs[0].add_run(str("1990-2021"))
    run1.font.name = "游明朝"
    run1.font.size = Pt(10)
    run1.font.underline = False
    run1.font.bold = False    

    # 対応する報告書データの14つ目の表の2行目3列目のテキストを置き換える（落雷)
    table = doc.tables[14]  # 8つ目の表を取得
    cell = table.cell(2, 2)  # 3行目3列目のセルを取得

    # 既存のテキストを置換
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""  # テキストをクリア

    run1 = cell.paragraphs[0].add_run(str("1990-2021"))
    run1.font.name = "游明朝"
    run1.font.size = Pt(10)
    run1.font.underline = False
    run1.font.bold = False    

    # 対応する報告書データの15つ目の表の2行目3列目のテキストを置き換える（積雪)
    table = doc.tables[15]  # 15つ目の表を取得
    cell = table.cell(2, 2)  # 3行目3列目のセルを取得

    # 既存のテキストを置換
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""  # テキストをクリア

    run1 = cell.paragraphs[0].add_run(str("1990-2021"))
    run1.font.name = "游明朝"
    run1.font.size = Pt(10)
    run1.font.underline = False
    run1.font.bold = False    
###############################################################################################################################################
###############################################################################################################################################


#最大瞬間風速#  

    # 気象庁データの12行目を読み込み（瞬間最大風速）
    max_wind_speed = str(meteorological_df.iloc[10, 0])
    
    # 対応する報告書データの2つ目の表の10行目4列目のテキストを置き換える
    table = doc.tables[2]  # 2つ目の表を取得
    cell = table.cell(9, 3)  # 10行目4列目のセルを取得

    # 既存のテキストを置換
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""  # テキストをクリア

    run1 = cell.paragraphs[0].add_run(str(max_wind_speed+"m/s"))
    run1.font.name = "游明朝"
    run1.font.size = Pt(10)
    run1.font.underline = False
    run1.font.bold = False

    # 対応する報告書データの8つ目の表の3行目4列目のテキストを置き換える
    table = doc.tables[8]  # 8つ目の表を取得
    cell = table.cell(2, 3)  # 3行目4列目のセルを取得

    # 既存のテキストを置換
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""  # テキストをクリア

    run1 = cell.paragraphs[0].add_run(str(max_wind_speed))
    run1.font.name = "游明朝"
    run1.font.size = Pt(10)
    run1.font.underline = False
    run1.font.bold = False    

    # 気象庁データの13行目を読み込み（瞬間最大風速観測年）
    max_wind_speed_year = str(meteorological_df.iloc[11, 0])
    
    # 対応する報告書データの2つ目の表の10行目5列目のテキストを置き換える
    table = doc.tables[2]  # 2つ目の表を取得
    cell = table.cell(9, 4)  # 10行目5列目のセルを取得

    # 既存のテキストを置換
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""  # テキストをクリア

    run1 = cell.paragraphs[0].add_run(str(max_wind_speed_year+"年"))
    run1.font.name = "游明朝"
    run1.font.size = Pt(10)
    run1.font.underline = False
    run1.font.bold = False

    # 対応する報告書データの8つ目の表の3行目5列目のテキストを置き換える
    table = doc.tables[8]  # 8つ目の表を取得
    cell = table.cell(2, 4)  # 3行目5列目のセルを取得

    # 既存のテキストを置換
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""  # テキストをクリア

    run1 = cell.paragraphs[0].add_run(str(max_wind_speed_year))
    run1.font.name = "游明朝"
    run1.font.size = Pt(10)
    run1.font.underline = False
    run1.font.bold = False 


    # 気象庁データの8行目を読み込み（瞬間最大風速観測点）
    max_wind_speed_point = str(meteorological_df.iloc[6, 0])
    
    # 対応する報告書データの2つ目の表の10行目5列目のテキストを置き換える
    table = doc.tables[2]  # 2つ目の表を取得
    cell = table.cell(9, 5)  # 10行目6列目のセルを取得

    # 既存のテキストを置換
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""  # テキストをクリア

    run1 = cell.paragraphs[0].add_run(str("(観測点："+max_wind_speed_point+")"))
    run1.font.name = "游明朝"
    run1.font.size = Pt(10)
    run1.font.underline = False
    run1.font.bold = False

   # 対応する報告書データの8つ目の表の2行目1列目のテキストを置き換える
    table = doc.tables[8]  # 8つ目の表を取得
    cell = table.cell(2, 0)  # 2行目1列目のセルを取得

    # 既存のテキストを置換
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""  # テキストをクリア

    run1 = cell.paragraphs[0].add_run(str(max_wind_speed_point))
    run1.font.name = "游明朝"
    run1.font.size = Pt(10)
    run1.font.underline = False
    run1.font.bold = False 

###############################################################################################################################################
###############################################################################################################################################

#最大風速#

    # 気象庁データの6行目を読み込み（最大風速）
    max_wind_speed2 = str(meteorological_df.iloc[4, 0])
    
    # 対応する報告書データの2つ目の表の10行目4列目のテキストを置き換える
    table = doc.tables[2]  # 2つ目の表を取得
    cell = table.cell(10, 3)  # 11行目4列目のセルを取得

    # 既存のテキストを置換
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""  # テキストをクリア
    # '所在地'の部分を追加

    run1 = cell.paragraphs[0].add_run(str(max_wind_speed2+"m/s"))
    run1.font.name = "游明朝"
    run1.font.size = Pt(10)
    run1.font.underline = False
    run1.font.bold = False

    # 対応する報告書データの8つ目の表の3行目4列目のテキストを置き換える
    table = doc.tables[8]  # 2つ目の表を取得
    cell = table.cell(3, 3)  # 4行目4列目のセルを取得

    # 既存のテキストを置換
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""  # テキストをクリア
    # '所在地'の部分を追加

    run1 = cell.paragraphs[0].add_run(str(max_wind_speed2))
    run1.font.name = "游明朝"
    run1.font.size = Pt(10)
    run1.font.underline = False
    run1.font.bold = False


    # 気象庁データの7行目を読み込み（最大風速観測年）
    max_wind_speed2_year = str(meteorological_df.iloc[5, 0])
    
    # 対応する報告書データの2つ目の表の10行目4列目のテキストを置き換える
    table = doc.tables[2]  # 2つ目の表を取得
    cell = table.cell(10, 4)  # 11行目5列目のセルを取得

    # 既存のテキストを置換
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""  # テキストをクリア
    # '所在地'の部分を追加

    run1 = cell.paragraphs[0].add_run(str(max_wind_speed2_year+"年"))
    run1.font.name = "游明朝"
    run1.font.size = Pt(10)
    run1.font.underline = False
    run1.font.bold = False 

    # 対応する報告書データの8つ目の表の10行目4列目のテキストを置き換える
    table = doc.tables[8]  # 2つ目の表を取得
    cell = table.cell(3, 4)  # 4行目5列目のセルを取得

    # 既存のテキストを置換
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""  # テキストをクリア
    # '所在地'の部分を追加

    run1 = cell.paragraphs[0].add_run(str(max_wind_speed2_year))
    run1.font.name = "游明朝"
    run1.font.size = Pt(10)
    run1.font.underline = False
    run1.font.bold = False 

    # 気象庁データの2行目を読み込み（最大風速観測点）
    max_wind_speed2_point = str(meteorological_df.iloc[0, 0])
    
    # 対応する報告書データの1つ目の表の11行目6列目のテキストを置き換える
    table = doc.tables[2]  # 2つ目の表を取得
    cell = table.cell(10, 5)  # 11行目6列目のセルを取得

    # 既存のテキストを置換
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""  # テキストをクリア

    run1 = cell.paragraphs[0].add_run(str("(観測点："+max_wind_speed_point+")"))
    run1.font.name = "游明朝"
    run1.font.size = Pt(10)
    run1.font.underline = False
    run1.font.bold = False

    # 対応する報告書データの9つ目の表の11行目6列目のテキストを置き換える
    table = doc.tables[9]  # 2つ目の表を取得
    cell = table.cell(2, 0)  # 2行目2列目のセルを取得

    # 既存のテキストを置換
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""  # テキストをクリア

    run1 = cell.paragraphs[0].add_run(str(max_wind_speed_point))
    run1.font.name = "游明朝"
    run1.font.size = Pt(10)
    run1.font.underline = False
    run1.font.bold = False
   
###############################################################################################################################################
###############################################################################################################################################
#最大風速15m以上の日数（最大値）

    # 気象庁データの18行目を読み込み（最大風速15m以上の日数）
    max_wind_speed_15mday = str(meteorological_df.iloc[16, 0])
    
    # 対応する報告書データの1つ目の表の13行目4列目のテキストを置き換える
    table = doc.tables[2]  # 2つ目の表を取得
    cell = table.cell(11, 3)  # 12行目4列目のセルを取得

    # 既存のテキストを置換
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""  # テキストをクリア

    run1 = cell.paragraphs[0].add_run(str(max_wind_speed_15mday+"日"))
    run1.font.name = "游明朝"
    run1.font.size = Pt(10)
    run1.font.underline = False
    run1.font.bold = False
    
    # 対応する報告書データの9つ目の表の3行目4列目のテキストを置き換える
    table = doc.tables[9]  # 9つ目の表を取得
    cell = table.cell(2, 3)  # 3行目4列目のセルを取得

    # 既存のテキストを置換
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""  # テキストをクリア

    run1 = cell.paragraphs[0].add_run(str(max_wind_speed_15mday))
    run1.font.name = "游明朝"
    run1.font.size = Pt(10)
    run1.font.underline = False
    run1.font.bold = False

    # 気象庁データの19行目を読み込み（最大風速15m以上の観測年）
    max_wind_speed_15mday_year = str(meteorological_df.iloc[17, 0])
    
    # 対応する報告書データの1つ目の表の19行目4列目のテキストを置き換える
    table = doc.tables[2]  # 2つ目の表を取得
    cell = table.cell(11, 4)  # 12行目5列目のセルを取得

    # 既存のテキストを置換
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""  # テキストをクリア

    run1 = cell.paragraphs[0].add_run(str(max_wind_speed_15mday_year+"年"))
    run1.font.name = "游明朝"
    run1.font.size = Pt(10)
    run1.font.underline = False
    run1.font.bold = False

    # 対応する報告書データの9つ目の表の3行目5列目のテキストを置き換える
    table = doc.tables[9]  # 2つ目の表を取得
    cell = table.cell(2, 4)  # 3行目5列目のセルを取得

    # 既存のテキストを置換
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""  # テキストをクリア

    run1 = cell.paragraphs[0].add_run(str(max_wind_speed_15mday_year))
    run1.font.name = "游明朝"
    run1.font.size = Pt(10)
    run1.font.underline = False
    run1.font.bold = False

    # 気象庁データの14行目を読み込み（最大風速15m以上の観測点）
    max_wind_speed_15mday_point = str(meteorological_df.iloc[12, 0])
    
    # 対応する報告書データの1つ目の表の19行目4列目のテキストを置き換える
    table = doc.tables[2]  # 2つ目の表を取得
    cell = table.cell(11, 5)  # 12行目6列目のセルを取得

    # 既存のテキストを置換
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""  # テキストをクリア

    run1 = cell.paragraphs[0].add_run(str("(観測点："+max_wind_speed_15mday_point+")"))
    run1.font.name = "游明朝"
    run1.font.size = Pt(10)
    run1.font.underline = False
    run1.font.bold = False


###############################################################################################################################################
###############################################################################################################################################
#落雷
    # 気象庁データの37行目を読み込み（年間平均落雷数）
    rakurai= str(meteorological_df.iloc[35, 0])
    
    # 対応する報告書データの1つ目の表の22行目3列目のテキストを置き換える
    table = doc.tables[2]  # 2つ目の表を取得
    cell = table.cell(20, 4)  # 22行目3列目のセルを取得

    # 既存のテキストを置換
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""  # テキストをクリア

    run1 = cell.paragraphs[0].add_run(str(max_wind_speed_15mday+"日"))
    run1.font.name = "游明朝"
    run1.font.size = Pt(10)
    run1.font.underline = False
    run1.font.bold = False

    # 対応する報告書データの14つ目の表の3行目8列目のテキストを置き換える
    table = doc.tables[14]  # 2つ目の表を取得
    cell = table.cell(2, 7)  # 3行目8列目のセルを取得

    # 既存のテキストを置換
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""  # テキストをクリア

    run1 = cell.paragraphs[0].add_run(str(max_wind_speed_15mday))
    run1.font.name = "游明朝"
    run1.font.size = Pt(10)
    run1.font.underline = False
    run1.font.bold = False

    # 気象庁データの29行目を読み込み（落雷観測点）
    rakurai_point= str(meteorological_df.iloc[27, 0])
    
    # 対応する報告書データの1つ目の表の22行目3列目のテキストを置き換える
    table = doc.tables[2]  # 2つ目の表を取得
    cell = table.cell(20, 5)  # 22行目4列目のセルを取得

    # 既存のテキストを置換
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""  # テキストをクリア

    run1 = cell.paragraphs[0].add_run(str(("(観測点："+rakurai_point+")")))
    run1.font.name = "游明朝"
    run1.font.size = Pt(10)
    run1.font.underline = False
    run1.font.bold = False

    # 対応する報告書データの14つ目の表の3行目1列目のテキストを置き換える
    table = doc.tables[14]  # 2つ目の表を取得
    cell = table.cell(2, 0)  # 3行目1列目のセルを取得

    # 既存のテキストを置換
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""  # テキストをクリア

    run1 = cell.paragraphs[0].add_run(str((rakurai_point)))
    run1.font.name = "游明朝"
    run1.font.size = Pt(10)
    run1.font.underline = False
    run1.font.bold = False

    # 気象庁データの33行目を読み込み（落雷最大値）
    rakurai_max= str(meteorological_df.iloc[31, 0])
    
    # 対応する報告書データの14つ目の表の3行目4列目のテキストを置き換える
    table = doc.tables[14]  # 2つ目の表を取得
    cell = table.cell(2, 3)  # 3行目4列目のセルを取得

    # 既存のテキストを置換
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""  # テキストをクリア

    run1 = cell.paragraphs[0].add_run(str((rakurai_max)))
    run1.font.name = "游明朝"
    run1.font.size = Pt(10)
    run1.font.underline = False
    run1.font.bold = False

    # 気象庁データの34行目を読み込み（落雷最大値年）
    rakurai_max_year= str(meteorological_df.iloc[32, 0])
    
    # 対応する報告書データの14つ目の表の3行目4列目のテキストを置き換える
    table = doc.tables[14]  # 2つ目の表を取得
    cell = table.cell(2, 4)  # 3行目4列目のセルを取得

    # 既存のテキストを置換
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""  # テキストをクリア

    run1 = cell.paragraphs[0].add_run(str((rakurai_max_year)))
    run1.font.name = "游明朝"
    run1.font.size = Pt(10)
    run1.font.underline = False
    run1.font.bold = False


    # 気象庁データの35行目を読み込み（落雷最小値）
    rakurai_min= str(meteorological_df.iloc[33, 0])
    
    # 対応する報告書データの14つ目の表の3行目4列目のテキストを置き換える
    table = doc.tables[14]  # 2つ目の表を取得
    cell = table.cell(2, 5)  # 3行目4列目のセルを取得

    # 既存のテキストを置換
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""  # テキストをクリア

    run1 = cell.paragraphs[0].add_run(str((rakurai_min)))
    run1.font.name = "游明朝"
    run1.font.size = Pt(10)
    run1.font.underline = False
    run1.font.bold = False

    # 気象庁データの36行目を読み込み（落雷最小値年）
    rakurai_min_year= str(meteorological_df.iloc[34, 0])
    
    # 対応する報告書データの14つ目の表の3行目4列目のテキストを置き換える
    table = doc.tables[14]  # 2つ目の表を取得
    cell = table.cell(2, 6)  # 3行目4列目のセルを取得

    # 既存のテキストを置換
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""  # テキストをクリア

    run1 = cell.paragraphs[0].add_run(str((rakurai_min_year)))
    run1.font.name = "游明朝"
    run1.font.size = Pt(10)
    run1.font.underline = False
    run1.font.bold = False

###############################################################################################################################################
###############################################################################################################################################
#積雪
    # 気象庁データの27行目を読み込み（最深積雪（最大値））
    sekisetsu_max= str(meteorological_df.iloc[25, 0])
    
    # 対応する報告書データの2つ目の表の23行目3列目のテキストを置き換える
    table = doc.tables[2]  # 2つ目の表を取得
    cell = table.cell(21, 3)  # 23行目3列目のセルを取得

    # 既存のテキストを置換
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""  # テキストをクリア

    run1 = cell.paragraphs[0].add_run(str(sekisetsu_max+"cm"))
    run1.font.name = "游明朝"
    run1.font.size = Pt(10)
    run1.font.underline = False
    run1.font.bold = False

    # 対応する報告書データの15つ目の表の3行目4列目のテキストを置き換える
    table = doc.tables[15]  # 15つ目の表を取得
    cell = table.cell(2, 3)  # 3行目4列目のセルを取得

    # 既存のテキストを置換
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""  # テキストをクリア

    run1 = cell.paragraphs[0].add_run(str(sekisetsu_max))
    run1.font.name = "游明朝"
    run1.font.size = Pt(10)
    run1.font.underline = False
    run1.font.bold = False


    # 気象庁データの28行目を読み込み（最深積雪（最大値）年）
    sekisetsu_max_year= str(meteorological_df.iloc[26, 0])
    
    # 対応する報告書データの2つ目の表の23行目3列目のテキストを置き換える
    table = doc.tables[2]  # 2つ目の表を取得
    cell = table.cell(21, 4)  # 23行目3列目のセルを取得

    # 既存のテキストを置換
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""  # テキストをクリア

    run1 = cell.paragraphs[0].add_run(str(sekisetsu_max_year+"年"))
    run1.font.name = "游明朝"
    run1.font.size = Pt(10)
    run1.font.underline = False
    run1.font.bold = False

    # 対応する報告書データの15つ目の表の23行目3列目のテキストを置き換える
    table = doc.tables[15]  # 2つ目の表を取得
    cell = table.cell(2, 4)  # 3行目5列目のセルを取得

    # 既存のテキストを置換
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""  # テキストをクリア

    run1 = cell.paragraphs[0].add_run(str(sekisetsu_max_year))
    run1.font.name = "游明朝"
    run1.font.size = Pt(10)
    run1.font.underline = False
    run1.font.bold = False

    # 気象庁データの23行目を読み込み（最深積雪（最大値）観測点）
    sekisetsu_max_point= str(meteorological_df.iloc[21, 0])
    
    # 対応する報告書データの2つ目の表の23行目3列目のテキストを置き換える
    table = doc.tables[2]  # 2つ目の表を取得
    cell = table.cell(21, 5)  # 23行目3列目のセルを取得

    # 既存のテキストを置換
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""  # テキストをクリア

    run1 = cell.paragraphs[0].add_run(str("(観測点："+sekisetsu_max_point+")"))
    run1.font.name = "游明朝"
    run1.font.size = Pt(10)
    run1.font.underline = False
    run1.font.bold = False

    # 対応する報告書データの15つ目の表の23行目3列目のテキストを置き換える
    table = doc.tables[15]  # 15つ目の表を取得
    cell = table.cell(2, 0)  # 3行目0列目のセルを取得

    # 既存のテキストを置換
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""  # テキストをクリア

    run1 = cell.paragraphs[0].add_run(str(sekisetsu_max_point))
    run1.font.name = "游明朝"
    run1.font.size = Pt(10)
    run1.font.underline = False
    run1.font.bold = False

###############################################################################################################################################
###############################################################################################################################################
#日照時間
    # 気象庁データの46行目を読み込み（日照時間）
    nisyotime = float(meteorological_df.iloc[44, 0])
    nisyotime = round(nisyotime, 1)
    nisyotime = str(nisyotime)

    
    # 対応する報告書データの2つ目の表の24行目3列目のテキストを置き換える
    table = doc.tables[2]  # 2つ目の表を取得
    cell = table.cell(22, 3)  # 23行目3列目のセルを取得

    # 既存のテキストを置換
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""  # テキストをクリア

    run1 = cell.paragraphs[0].add_run(str(nisyotime+"時間"))
    run1.font.name = "游明朝"
    run1.font.size = Pt(10)
    run1.font.underline = False
    run1.font.bold = False






    # 書き換えた報告書データを新しいファイルに保存
    report_data_file = os.path.join(output_folder, f'{report_data_name}')
    doc.save(report_data_file)



print("処理が完了しました。")
